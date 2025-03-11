from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
import os
from docx import Document
from groq import Groq
import tempfile
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from datetime import datetime
from docx.oxml import parse_xml
from docx.shared import Inches
import re
import time

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def extract_formatting(doc_path):
    doc = Document(doc_path)
    formatting = {
        'sections': [],
        'styles': {}
    }
    for section in doc.sections:
        section_format = {
            'page_height': section.page_height,
            'page_width': section.page_width,
            'left_margin': section.left_margin,
            'right_margin': section.right_margin,
            'top_margin': section.top_margin,
            'bottom_margin': section.bottom_margin
        }
        formatting['sections'].append(section_format)
    
    return formatting

def calculate_chapter_distribution(num_pages):
    """Calculate detailed word counts for each section based on pages"""
    total_words = num_pages * 300
    
    distribution = {
        1: {  # Introduction (15%)
            'total_words': int(total_words * 0.15),
            'sections': {
                '1.1': 0.25,  # Identification of Client/Need
                '1.2': 0.20,  # Identification of Problem
                '1.3': 0.20,  # Identification of Tasks
                '1.4': 0.15,  # Timeline
                '1.5': 0.20   # Organization of Report
            }
        },
        2: {  # Literature Review (25%)
            'total_words': int(total_words * 0.25),
            'sections': {
                '2.1': 0.15,  # Timeline of problem
                '2.2': 0.25,  # Existing solutions
                '2.3': 0.20,  # Bibliometric analysis
                '2.4': 0.15,  # Review Summary
                '2.5': 0.15,  # Problem Definition
                '2.6': 0.10   # Goals/Objectives
            }
        },
        3: {  # Design Flow (25%)
            'total_words': int(total_words * 0.25),
            'sections': {
                '3.1': 0.20,  # Evaluation & Selection
                '3.2': 0.20,  # Design Constraints
                '3.3': 0.20,  # Analysis of Features
                '3.4': 0.20,  # Design Flow
                '3.5': 0.20   # Design selection
            }
        },
        4: {  # Results (20%)
            'total_words': int(total_words * 0.20),
            'sections': {
                '4.1': 0.35,  # Implementation details
                '4.2': 0.35,  # Results analysis
                '4.3': 0.30   # Validation
            }
        },
        5: {  # Conclusion (15%)
            'total_words': int(total_words * 0.15),
            'sections': {
                '5.1': 0.60,  # Conclusion
                '5.2': 0.40   # Future work
            }
        }
    }
    return distribution

def generate_section_content(title, chapter_num, section_num, target_words, context=""):
    """Generate content for a specific section with word count control"""
    client = Groq()
    
    prompt = f"""Generate section {chapter_num}.{section_num} for "{title}".
    Structure the content as follows:
    • Use only two-level section numbering ({chapter_num}.{section_num})
    • Use exactly two asterisks (**) for bold text, not four asterisks (****)
    • Format headings as: **{chapter_num}.{section_num} Title**
    • Format subheadings as: **Subheading Title**
    • For bullet points:
      - Start with • 
      - Use **Key Term:** for emphasized terms
    • Example format:
      **1.1 Overview**
      [Introduction paragraph]
      
      **System Architecture**
      • **Database:** Description...
      • **Network:** Description...
      
    Target length: {target_words} words.
    Previous context: {context}"""
    
    # Add delay between API calls to prevent rate limiting
    time.sleep(2)
    
    completion = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {
                "role": "system",
                "content": "Generate detailed academic content for a technical project report section. Maintain consistent formatting and technical depth."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.7,
        max_tokens=2048
    )
    
    return process_content_section(completion.choices[0].message.content)

def generate_project_report(title, num_pages, formatting):
    """Generate report content in carefully controlled chunks"""
    distribution = calculate_chapter_distribution(num_pages)
    content = []
    context = ""
    
    # Generate content chapter by chapter, section by section
    for chapter_num in range(1, 6):
        chapter_content = []
        chapter_info = distribution[chapter_num]
        
        # Add chapter heading
        chapter_titles = {
            1: "CHAPTER 1. INTRODUCTION",
            2: "CHAPTER 2. LITERATURE REVIEW/BACKGROUND STUDY",
            3: "CHAPTER 3. DESIGN FLOW/PROCESS",
            4: "CHAPTER 4. RESULTS ANALYSIS AND VALIDATION",
            5: "CHAPTER 5. CONCLUSION AND FUTURE WORK"
        }
        chapter_content.append(chapter_titles[chapter_num])
        
        # Generate each section with proper sequential numbering
        section_keys = sorted(chapter_info['sections'].keys(), key=lambda x: float(x))
        for section_key in section_keys:
            proportion = chapter_info['sections'][section_key]
            section_words = int(chapter_info['total_words'] * proportion)
            
            # Extract section number from the key (e.g., '1.1' -> 1)
            section_num = section_key.split('.')[1]
            
            section_content = generate_section_content(
                title,
                chapter_num,
                section_num,
                section_words,
                context
            )
            chapter_content.append(section_content)
            # Update context for next section
            context = f"{context}\n{section_content}"[-500:]  # Keep last 500 chars for context
        
        content.append("\n\n".join(chapter_content))
        # Add delay between chapters
        time.sleep(5)
    
    # Generate references
    references = generate_references(title)
    
    # Combine all content
    full_content = "\n\n".join(content) + "\n\nREFERENCES\n" + references
    
    return full_content

def generate_references(title):
    """Generate IEEE formatted references relevant to the project topic"""
    client = Groq()
    
    prompt = f"""Generate 15-20 relevant academic references for a project report about "{title}".
    Requirements:
    1. Use IEEE citation format
    2. Include recent papers (last 5-10 years)
    3. Focus on reputable journals and conferences
    4. Mix of foundational and recent works
    5. Ensure relevance to {title}
    6. Do not include any introductory text
    7. Start directly with numbered references
    8. Each reference should be on a new line
    9. Format exactly like this:
       [1] A. Author, B. Author and C. Author, "Title of paper," Name of Journal, vol. x, no. x, pp. xxx-xxx, Month Year.
    """
    
    # Add delay to prevent rate limiting
    time.sleep(2)
    
    completion = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {
                "role": "system",
                "content": "Generate academic references in IEEE format. Start directly with the numbered references. Do not include any introductory text."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.7,
        max_tokens=2048
    )
    
    # Process and clean up references
    references = completion.choices[0].message.content
    
    # Remove any introductory text
    references = re.sub(r'^.*?(?=\[1\])', '', references, flags=re.DOTALL)
    
    # Clean up formatting
    references = re.sub(r'\n{3,}', '\n\n', references)  # Remove extra newlines
    references = re.sub(r'^\s+', '', references, flags=re.MULTILINE)  # Remove leading spaces
    references = references.strip()
    
    # Add proper spacing between references
    references = re.sub(r'(\[\d+\])', r'\n\1', references)
    references = references.strip()
    
    # Add proper heading
    references = "REFERENCES\n\n" + references
    
    return references

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_report():
    title = request.form['title']
    num_pages = int(request.form['num_pages'])
    current_month_year = datetime.now().strftime("%b %Y")
    
    try:
        # Create new document
        doc = Document()
        
        # Title Page
        # "A PROJECT REPORT"
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("A PROJECT REPORT")
        title_run.font.size = Pt(18)
        title_run.font.name = 'Times New Roman'
        title_run.bold = True
        title_para.paragraph_format.line_spacing = 1.5
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().add_run().add_break()
        
        # "Submitted by"
        submitted_para = doc.add_paragraph()
        submitted_run = submitted_para.add_run("Submitted by")
        submitted_run.font.size = Pt(14)
        submitted_run.font.name = 'Times New Roman'
        submitted_run.bold = True
        submitted_run.italic = True
        submitted_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().add_run().add_break()
        
        # Candidate Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run("[NAME OF THE CANDIDATE(S)]")
        name_run.font.size = Pt(16)
        name_run.font.name = 'Times New Roman'
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().add_run().add_break()
        
        # Degree fulfillment text
        degree_para = doc.add_paragraph()
        degree_run = degree_para.add_run("in partial fulfillment for the award of the degree of")
        degree_run.font.size = Pt(14)
        degree_run.font.name = 'Times New Roman'
        degree_run.bold = True
        degree_run.italic = True
        degree_para.paragraph_format.line_spacing = 1.5
        degree_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().add_run().add_break()
        
        # Degree Name
        degree_name_para = doc.add_paragraph()
        degree_name_run = degree_name_para.add_run("[NAME OF THE DEGREE]")
        degree_name_run.font.size = Pt(16)
        degree_name_run.font.name = 'Times New Roman'
        degree_name_run.bold = True
        degree_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Branch
        branch_para = doc.add_paragraph()
        branch_para.add_run("IN\n").font.size = Pt(14)
        branch_run = branch_para.add_run("[BRANCH OF STUDY]")
        branch_run.font.size = Pt(14)
        branch_run.font.name = 'Times New Roman'
        branch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().add_run().add_break()

        # Add logo
        logo_para = doc.add_paragraph()
        logo_run = logo_para.add_run()
        logo_run.add_picture('static/cu_logo.png', width=Pt(200))
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # University and Date
        univ_para = doc.add_paragraph()
        univ_run = univ_para.add_run("Chandigarh University")
        univ_run.font.size = Pt(14)
        univ_run.font.name = 'Times New Roman'
        univ_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(current_month_year)
        date_run.font.size = Pt(14)
        date_run.font.name = 'Times New Roman'
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break before certificate
        doc.add_page_break()
        
        # Add logo before Bonafide Certificate (same width as first page)
        logo_para = doc.add_paragraph()
        logo_run = logo_para.add_run()
        logo_run.add_picture('static/cu_logo.png', width=Pt(200))
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().add_run().add_break()
        # Bonafide Certificate
        cert_title = doc.add_paragraph()
        cert_run = cert_title.add_run("BONAFIDE CERTIFICATE")
        cert_run.font.size = Pt(16)
        cert_run.font.name = 'Times New Roman'
        cert_run.bold = True
        cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().add_run().add_break()
        
        # Certificate content with proper formatting
        cert_content = doc.add_paragraph()
        cert_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Add justification
        cert_text = f'Certified that this project report "{title}" is the '
        cert_run = cert_content.add_run(cert_text)
        cert_run.font.size = Pt(14)
        cert_run.font.name = 'Times New Roman'
        
        # Add title in bold
        title_run = cert_content.add_run(title)
        title_run.font.size = Pt(14)
        title_run.font.name = 'Times New Roman'
        title_run.bold = True
        
        # Continue text
        next_text = '" is the '
        next_run = cert_content.add_run(next_text)
        next_run.font.size = Pt(14)
        next_run.font.name = 'Times New Roman'
        
        # Add underlined 'bonafide' word
        bonafide_run = cert_content.add_run("bonafide")
        bonafide_run.font.size = Pt(14)
        bonafide_run.font.name = 'Times New Roman'
        bonafide_run.underline = True
        
        # Continue text
        mid_text = ' work of "'
        mid_run = cert_content.add_run(mid_text)
        mid_run.font.size = Pt(14)
        mid_run.font.name = 'Times New Roman'
        
        # Add candidate name in bold
        name_run = cert_content.add_run("[NAME OF THE CANDIDATE(S)]")
        name_run.font.size = Pt(14)
        name_run.font.name = 'Times New Roman'
        name_run.bold = True
        
        # Final text
        final_text = '" who carried out the project work under my/our supervision.'
        final_run = cert_content.add_run(final_text)
        final_run.font.size = Pt(14)
        final_run.font.name = 'Times New Roman'
        
        doc.add_paragraph().add_run().add_break()
        doc.add_paragraph().add_run().add_break()
        
        # Add signatures using tab stops
        signature_para = doc.add_paragraph()
        signature_para.paragraph_format.tab_stops.add_tab_stop(Inches(4.5))
        
        # Add SIGNATURE text in bold
        signature_text = "SIGNATURE\tSIGNATURE"
        signature_run = signature_para.add_run(signature_text)
        signature_run.font.name = 'Times New Roman'
        signature_run.font.size = Pt(12)
        signature_run.bold = True
        
        # Add underline for signatures
        signature_line = doc.add_paragraph()
        signature_line.paragraph_format.tab_stops.add_tab_stop(Inches(4.5))
        signature_line.add_run("_____________________\t_____________________")
        
        # Add designation in bold
        designation_para = doc.add_paragraph()
        designation_para.paragraph_format.tab_stops.add_tab_stop(Inches(4.5))
        designation_text = "HEAD OF THE DEPARTMENT\tSUPERVISOR"
        designation_run = designation_para.add_run(designation_text)
        designation_run.font.name = 'Times New Roman'
        designation_run.font.size = Pt(12)
        designation_run.bold = True
        
        doc.add_paragraph().add_run().add_break()
        doc.add_paragraph().add_run().add_break()
        
        # Add viva-voce text
        viva_para = doc.add_paragraph()
        viva_text = "Submitted for the project "
        viva_run = viva_para.add_run(viva_text)
        viva_run.font.size = Pt(12)
        viva_run.font.name = 'Times New Roman'
        
        # Add underlined 'viva-voce'
        viva_underline = viva_para.add_run("viva-voce")
        viva_underline.font.size = Pt(12)
        viva_underline.font.name = 'Times New Roman'
        viva_underline.underline = True
        
        # Add remaining text
        viva_end = viva_para.add_run(" examination held on _________________")
        viva_end.font.size = Pt(12)
        viva_end.font.name = 'Times New Roman'
        
        doc.add_paragraph().add_run().add_break()

        
        # Add examiners using tab stops
        examiner_para = doc.add_paragraph()
        examiner_para.paragraph_format.tab_stops.add_tab_stop(Inches(4))
        
        # Add underline for examiners
        examiner_line = doc.add_paragraph()
        examiner_line.paragraph_format.tab_stops.add_tab_stop(Inches(4))
        examiner_line.add_run("_____________________\t_____________________")
        
        # Add examiner text in bold
        examiner_text = "INTERNAL EXAMINER\tEXTERNAL EXAMINER"
        examiner_run = examiner_para.add_run(examiner_text)
        examiner_run.font.name = 'Times New Roman'
        examiner_run.font.size = Pt(12)
        examiner_run.bold = True
        
        # Add Table of Contents
        doc.add_page_break()
        
        # Add Table of Contents heading (16pt)
        toc_heading = doc.add_paragraph()
        toc_run = toc_heading.add_run("TABLE OF CONTENTS")
        toc_run.font.size = Pt(16)  # Main TOC heading is 16pt
        toc_run.font.name = 'Times New Roman'
        toc_run.bold = True
        toc_run.underline = True
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add space after heading
        
        # Add Lists with proper spacing and tab stops (14pt)
        for list_title, page_num in [
            ("List of Figures", 7),
            ("List of Tables", 8),
            ("List of Standards", 9)
        ]:
            para = doc.add_paragraph()
            para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            # Main list titles are 14pt
            list_run = para.add_run(list_title)
            list_run.font.name = 'Times New Roman'
            list_run.font.size = Pt(12)
            para.add_run(f'\t{page_num}')
        
        doc.add_paragraph()  # Add space before chapters
        
        # Add chapters and their sections
        chapters = {
            "CHAPTER 1. INTRODUCTION": {
                "sections": [
                    "1.1. Identification of Client/Need/ Relevant Contemporary issue",
                    "1.2. Identification of Problem",
                    "1.3. Identification of Tasks",
                    "1.4. Timeline",
                    "1.5. Organization of the Report"
                ],
                "page": 11
            },
            "CHAPTER 2. LITERATURE REVIEW/BACKGROUND STUDY": {
                "sections": [
                    "2.1. Timeline of the reported problem",
                    "2.2. Existing solutions",
                    "2.3. Bibliometric analysis",
                    "2.4. Review Summary",
                    "2.5. Problem Definition",
                    "2.6. Goals/Objectives"
                ],
                "page": 12
            },
            "CHAPTER 3. DESIGN FLOW/PROCESS": {
                "sections": [
                    "3.1. Evaluation & Selection of Specifications/Features",
                    "3.2. Design Constraints",
                    "3.3. Analysis of Features and finalization subject to constraints",
                    "3.4. Design Flow",
                    "3.5. Design selection",
                    "3.6. Implementation plan methodology"
                ],
                "page": 13
            },
            "CHAPTER 4. RESULTS ANALYSIS AND VALIDATION": {
                "sections": [
                    "4.1. Implementation of solution"
                ],
                "page": 14
            },
            "CHAPTER 5. CONCLUSION AND FUTURE WORK": {
                "sections": [
                    "5.1. Conclusion",
                    "5.2. Future work"
                ],
                "page": 15
            }
        }
        
        # Add chapters with proper tab stops
        for chapter, details in chapters.items():
            # Add chapter heading (14pt)
            chapter_para = doc.add_paragraph()
            chapter_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            chapter_run = chapter_para.add_run(chapter)
            chapter_run.font.name = 'Times New Roman'
            chapter_run.font.size = Pt(14)  # Chapter headings are 14pt
            chapter_run.bold = True
            chapter_para.add_run(f'\t{details["page"]}')

            # Add sections with proper indentation and tab stops (12pt)
            for section in details["sections"]:
                section_para = doc.add_paragraph()
                section_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                
                section_run = section_para.add_run(section)
                section_run.font.name = 'Times New Roman'
                section_run.font.size = Pt(12)  # Subheadings are 12pt
                section_para.add_run(f'\t{details["page"]}')

            doc.add_paragraph()  # Add space between chapters
        
        # Add final sections (REFERENCES, APPENDIX, USER MANUAL)
        final_sections = [
            ("REFERENCES", 16, []),
            ("APPENDIX", 17, [
                "1. Plagiarism Report",
                "2. Design Checklist"
            ]),
            ("USER MANUAL", 18, [])
        ]
        
        for section, page_num, subsections in final_sections:
            # Add main section (14pt)
            section_para = doc.add_paragraph()
            section_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            section_run = section_para.add_run(section)
            section_run.font.name = 'Times New Roman'
            section_run.font.size = Pt(14)  # Main sections are 14pt
            section_run.bold = True
            section_para.add_run(f'\t{page_num}')
            
            # Add subsections if any (12pt)
            for subsection in subsections:
                subsection_para = doc.add_paragraph()
                subsection_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                
                subsection_run = subsection_para.add_run(subsection)
                subsection_run.font.name = 'Times New Roman'
                subsection_run.font.size = Pt(12)  # Subsections are 12pt
                subsection_para.add_run(f'\t{page_num}')

            doc.add_paragraph()  # Add space after each main section
        
        # Remove the page break and directly start processing content
        # Generate content using AI
        content = generate_project_report(title, num_pages, {})
        
        # Process and add the content with proper formatting
        sections = content.split('\n\n')
        current_chapter = []
        
        # Skip any introductory text
        start_index = 0
        for i, section in enumerate(sections):
            if section.strip().startswith('CHAPTER'):
                start_index = i
                break
        
        # Add page break before first chapter
        doc.add_page_break()
        
        # Process all sections including the last chapter and references
        for i, section in enumerate(sections[start_index:]):
            if section.strip():
                # If we encounter a new chapter or references
                if section.strip().startswith('CHAPTER') or section.strip().startswith('REFERENCES'):
                    # If we have content from previous chapter, add it first
                    if current_chapter:
                        # Add all content from current chapter
                        for chapter_section in current_chapter:
                            # Handle chapter headings
                            if chapter_section.startswith('CHAPTER'):
                                para = doc.add_paragraph()
                                run = para.add_run(chapter_section)
                                run.font.size = Pt(16)
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                para.paragraph_format.line_spacing = 1.5
                            else:
                                lines = chapter_section.split('\n')
                                for line in lines:
                                    line = line.strip()
                                    if line:
                                        para = add_formatted_content(doc, line)
                                        para.paragraph_format.line_spacing = 1.5
                        
                        # Add page break after completing the chapter (except for references)
                        if not current_chapter[0].strip().startswith('REFERENCES'):
                            doc.add_page_break()
                        current_chapter = []
                    
                    # Start new chapter or references
                    current_chapter.append(section)
                else:
                    # Add section to current chapter
                    current_chapter.append(section)
        
        # Process the last chapter/section if any content remains
        if current_chapter:
            for chapter_section in current_chapter:
                if chapter_section.strip().startswith('REFERENCES'):
                    para = doc.add_paragraph()
                    run = para.add_run("REFERENCES")
                    run.font.size = Pt(16)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.line_spacing = 1.5
                    
                    # Add references content with proper formatting
                    ref_content = chapter_section.replace("REFERENCES", "").strip()
                    if ref_content:
                        # Split references into individual entries
                        ref_entries = re.split(r'(\[\d+\])', ref_content)
                        for entry in ref_entries:
                            if entry.strip():
                                ref_para = doc.add_paragraph()
                                ref_run = ref_para.add_run(entry.strip())
                                ref_run.font.size = Pt(12)
                                ref_run.font.name = 'Times New Roman'
                                ref_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                ref_para.paragraph_format.line_spacing = 1.5
                                ref_para.paragraph_format.left_indent = Inches(0.5)  # Add left indentation
                                ref_para.paragraph_format.first_line_indent = Inches(-0.5)  # Hanging indent
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(chapter_section.strip())
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.line_spacing = 1.5

        # Save the document after all content has been processed
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{secure_filename(title)}.docx")
        doc.save(output_path)
        
        return send_file(output_path, as_attachment=True)
        
    except Exception as e:
        return str(e), 500

def process_content_section(content):
    """Clean and format section content to maintain consistent styling"""
    processed_lines = []
    
    for line in content.split('\n'):
        # Skip empty lines
        if not line or len(line.strip()) == 0:
            processed_lines.append('')
            continue
            
        # Remove all asterisks from section numbers and subheadings
        if re.match(r'^[\*]+\d+\.\d+\s+', line) or re.match(r'^\d+\.\d+\s+', line):
            line = re.sub(r'^[\*]+|\*+$', '', line)
            processed_lines.append(line)
            continue
            
        # Remove asterisks from subheadings (including those with "and", "&", etc.)
        if (re.match(r'^[\*]+[A-Z][A-Za-z\s]+(and|&)?[A-Za-z\s]+', line) or 
            (len(line.split()) <= 4 and len(line.lstrip('*')) > 0 and line.lstrip('*')[0].isupper())):
            line = re.sub(r'^[\*]+|\*+$', '', line)
            processed_lines.append(line)
            continue
            
        # Handle bullet points and clean up asterisks after colons
        if line.strip().startswith('•'):
            # Remove extra asterisks after colons
            line = re.sub(r':\s*\*+', ':', line)
            # Clean up any remaining multiple asterisks
            line = re.sub(r'\*{2,}', '', line)
            processed_lines.append(line)
            continue
            
        # For all other lines, remove all asterisks
        line = re.sub(r'\*+', '', line)
        processed_lines.append(line)
    
    content = '\n'.join(processed_lines)
    content = content.strip()
    return content

def add_formatted_content(doc, line, content_para=None):
    """Helper function to handle text formatting"""
    if content_para is None:
        content_para = doc.add_paragraph()
    
    # Skip empty lines
    if not line or len(line.strip()) == 0:
        return content_para
    
    # Handle section headings (e.g., "1.2 Modernization...")
    if re.match(r'^\d+\.\d+\s+', line):
        run = content_para.add_run(line)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content_para.paragraph_format.space_before = Pt(12)
        content_para.paragraph_format.space_after = Pt(6)
        return content_para

    # Handle subheadings (improved pattern to catch more cases)
    if ((re.match(r'^[A-Z][A-Za-z\s]+(and|&)?[A-Za-z\s]+', line) and len(line.split()) <= 6) or
        (len(line.split()) <= 4 and len(line) > 0 and line[0].isupper())):
        run = content_para.add_run(line)
        run.bold = True
        run.font.size = Pt(12)  # Changed to 14pt
        run.font.name = 'Times New Roman'
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content_para.paragraph_format.space_before = Pt(12)
        content_para.paragraph_format.space_after = Pt(6)
        return content_para

    # Handle bullet points
    if line.startswith('•'):
        line = line.replace('•', '').strip()
        content_para.style = 'List Bullet'
        content_para.paragraph_format.left_indent = Inches(0.5)
        
        # Handle bullet points with colons
        if ':' in line:
            before_colon, after_colon = line.split(':', 1)
            run = content_para.add_run(before_colon + ':')
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            
            run = content_para.add_run(after_colon)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        else:
            run = content_para.add_run(line)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return content_para

    # Handle regular text
    run = content_para.add_run(line)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return content_para

if __name__ == '__main__':
    app.run(debug=True) 